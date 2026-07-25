import streamlit as st
import google.generativeai as genai
import csv
import io
import pandas as pd
import re
import random
from docx import Document

# ==========================================
# 1. ページ設定 & カスタムCSS
# ==========================================
st.set_page_config(
    page_title="ツナグ先生 - 統合校務支援システム (V9.4大容量データ版)",
    page_icon="🏫",
    layout="wide"
)

# 💡 上部余白調整とグループ見出しの文字欠け防止CSS
st.markdown("""
    <style>
    /* 上部余白を適正値に調整（見切れるのを防止） */
    .block-container {
        padding-top: 2.5rem !important;
        padding-bottom: 1rem !important;
    }
    
    .student-card { 
        background-color: #f8f9fa; 
        border: 1px solid #dee2e6; 
        padding: 15px; 
        border-radius: 8px; 
        margin-bottom: 15px; 
    }

    /* グループ見出しのスタイル（line-heightとpaddingで高さ・文字表示を最適化） */
    .menu-group-header {
        font-size: 0.95rem;
        font-weight: bold;
        line-height: 1.5 !important;
        margin-top: 8px;
        margin-bottom: 6px;
        padding: 6px 10px;
        border-radius: 4px;
        display: block;
    }
    .homeroom-header {
        color: #d9534f;
        background-color: #fdf2f2;
        border-left: 4px solid #d9534f;
    }
    .subject-header {
        color: #1f77b4;
        background-color: #f0f7fc;
        border-left: 4px solid #1f77b4;
    }

    /* 1段目：学級担任用ボタンの色付け（オレンジ系） */
    div[data-testid="stSegmentedControl"]:nth-of-type(1) button[aria-selected="true"] {
        background-color: #e67e22 !important;
        color: white !important;
    }

    /* 2段目：教科担当用ボタンの色付け（ブルー系） */
    div[data-testid="stSegmentedControl"]:nth-of-type(2) button[aria-selected="true"] {
        background-color: #1f77b4 !important;
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 2. 大容量ダミーデータ生成ロジック
# ==========================================

@st.cache_data
def generate_full_dummy_data():
    surnames = ["佐藤", "鈴木", "高橋", "田中", "伊藤", "渡辺", "山本", "中村", "小林", "加藤", 
                "吉田", "山田", "佐々木", "山口", "松本", "井上", "木村", "林", "斎藤", "清水"]
    male_names = ["蓮", "悠真", "湊", "大翔", "樹", "陽翔", "悠人", "颯太", "陸", "翔太"]
    female_names = ["葵", "陽葵", "凛", "結菜", "芽依", "詩", "結愛", "莉子", "咲良", "結衣"]

    classes_config = [
        ("1年1組", 40),
        ("1年5組", 40),
        ("2年2組", 40),
        ("3年3組", 40),
        ("3年5組", 40)
    ]

    master_list = []
    score_list = []
    
    random.seed(42)

    for cls_name, count in classes_config:
        for num in range(1, count + 1):
            gender = "男" if num % 2 != 0 else "女"
            sname = surnames[(num - 1) % len(surnames)]
            gname = male_names[(num - 1) % len(male_names)] if gender == "男" else female_names[(num - 1) % len(female_names)]
            full_name = f"{sname} {gname}"

            master_list.append({
                "クラス": cls_name,
                "出席番号": num,
                "氏名": full_name,
                "性別": gender,
                "ステータス": "在籍",
                "異動日": "",
                "備考": "担任クラス" if cls_name == "1年1組" else "授業担当"
            })

            mid_score = random.randint(45, 98)
            final_score = random.randint(50, 100)
            k1 = random.randint(55, 98)
            k2 = random.randint(50, 95)
            k3 = random.randint(60, 100)
            avg = (k1 + k2 + k3) / 3
            grade = 5 if avg >= 85 else (4 if avg >= 70 else (3 if avg >= 55 else (2 if avg >= 40 else 1)))

            score_list.append({
                "クラス": cls_name,
                "出席番号": num,
                "氏名": full_name,
                "中間テスト": mid_score,
                "期末テスト": final_score,
                "観点1_知識": k1,
                "観点2_思考": k2,
                "観点3_主体性": k3,
                "評定": grade,
                "総合所見": f"{'課題に対して粘り強く思考し、工夫して解決策を導くことができました。' if grade >= 4 else '基礎的な計算力の定着が見られ、授業中の挙手・発言も意欲的です。'}"
            })

    logs_list = []
    dates = ["2026-04-15", "2026-04-22", "2026-05-10", "2026-05-18", "2026-06-04", "2026-06-15", "2026-07-02"]
    memo_templates = [
        ("数学", "方程式の立式において、自力で関係性を見つけ出して正解を導き出すことができた。"),
        ("数学", "グループワークで解き方に悩んでいる班員に対して丁寧にやり方を教えていた。"),
        ("総合・行動", "行事の実行委員に立候補し、クラス全体の意見をスムーズに集約・発表した。"),
        ("総合・行動", "清掃活動において自分の担当場所が終わった後、進んで共有スペースの掃除を手伝った。"),
        ("特別活動", "朝の読書時間に毎日落ち着いて読書に取り組み、クラスの静寂な雰囲気作りに貢献した。"),
        ("国語・他", "朝の会での1分スピーチにて、自分の体験に基づいた説得力のある発表を行った。"),
        ("生活指導", "係活動のプリント配布を毎朝忘れずに行い、責任感を持って行動できている。")
    ]

    c1_students = [m for m in master_list if m["クラス"] == "1年1組"]
    for idx, st_item in enumerate(c1_students):
        num_memos = 2 if idx % 2 == 0 else 3
        for i in range(num_memos):
            d = dates[(idx + i) % len(dates)]
            cat, memo_text = memo_templates[(idx * 2 + i) % len(memo_templates)]
            logs_list.append({
                "日付": d,
                "クラス": "1年1組",
                "出席番号": st_item["出席番号"],
                "氏名": st_item["氏名"],
                "対象分野": cat,
                "観察メモ": memo_text
            })

    return pd.DataFrame(master_list), pd.DataFrame(score_list), pd.DataFrame(logs_list)


# ==========================================
# 3. データベース（セッション状態）の初期化
# ==========================================

if "school_year" not in st.session_state:
    st.session_state.school_year = "2026"
if "teacher_name" not in st.session_state:
    st.session_state.teacher_name = "山田 太郎"
if "active_menu" not in st.session_state:
    st.session_state.active_menu = "📝 ① 日々メモ蓄積 & 所見"

df_master, df_scores, df_logs = generate_full_dummy_data()

if "student_master" not in st.session_state:
    st.session_state.student_master = df_master

if "daily_logs" not in st.session_state:
    st.session_state.daily_logs = df_logs

if "subject_scores" not in st.session_state:
    st.session_state.subject_scores = df_scores

api_key = ""
try:
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
except Exception:
    pass

# --- サイドバーエリア ---
with st.sidebar:
    # 💡 パターン1：タイトルを左側ペインに移動
    st.title("🏫 ツナグ先生")
    st.caption(f"統合校務支援システム (V9.4 / {st.session_state.school_year}年度)")
    st.markdown("---")

    st.header("⚙️ システム・基本情報")
    st.session_state.school_year = st.text_input("年度設定", st.session_state.school_year)
    st.session_state.teacher_name = st.text_input("担任教員名", st.session_state.teacher_name)
    
    st.markdown("---")
    st.info("🏫 **担任:** 1年1組 (40名)\n📚 **担当:** 1-1, 1-5, 2-2, 3-3, 3-5 (計200名)")
    st.markdown("---")
    
    if not api_key:
        api_key = st.text_input("Gemini API Key", type="password")
    else:
        st.success("🔑 API連携中")
    if api_key:
        genai.configure(api_key=api_key)
        
    doc_type = st.radio("文末モード:", ["です・ます調（通知表）", "である・した調（要録）"])
    ending_rule = "文末は「です・ます」調で統一。" if "です" in doc_type else "文末は「である・した」調で統一。"

# Helper Functions
def replace_docx_tags(doc, data_dict):
    for p in doc.paragraphs:
        for key, value in data_dict.items():
            tag = f"{{{{{key}}}}}"
            if tag in p.text:
                p.text = p.text.replace(tag, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in data_dict.items():
                        tag = f"{{{{{key}}}}}"
                        if tag in p.text:
                            p.text = p.text.replace(tag, str(value))
    return doc

def get_all_classes():
    classes = sorted(st.session_state.student_master["クラス"].dropna().unique().tolist())
    return classes if classes else ["1年1組"]

def get_active_students(cls_name=None):
    df = st.session_state.student_master
    active_df = df[df["ステータス"] == "在籍"]
    if cls_name:
        active_df = active_df[active_df["クラス"] == cls_name]
    return active_df["氏名"].tolist()


# ==========================================
# 4. メイン画面（担任/教科 視覚的グループ分けメニュー）
# ==========================================

# 💡 学級担任用・教科担当用でメニューオプションを分割
homeroom_options = [
    "📝 ① 日々メモ蓄積 & 所見",
    "📁 ② CSV一括生成",
    "🔍 ③ 所見データ自動校正",
    "💬 ④ 蓄積連動カルテ",
    "🔄 ⑦ 担任用 全教科成績集約",
    "🖨️ ⑧ 通知表・要録 印刷＆出力"
]

subject_options = [
    "📊 ⑤ 成績＆評定自動計算",
    "📈 ⑥ 学期推移ダッシュボード",
    "⚙️ ⓪ 担任＆授業担当 名簿管理"
]

# メニュー選択の同期処理
def sync_menu_homeroom():
    st.session_state.active_menu = st.session_state.hr_menu

def sync_menu_subject():
    st.session_state.active_menu = st.session_state.sub_menu

# --- 1段目：学級担任向け機能 ---
st.markdown('<div class="menu-group-header homeroom-header">🏠 学級担任メイン機能</div>', unsafe_allow_html=True)
st.segmented_control(
    "学級担任機能",
    options=homeroom_options,
    default=st.session_state.active_menu if st.session_state.active_menu in homeroom_options else None,
    key="hr_menu",
    on_change=sync_menu_homeroom,
    label_visibility="collapsed"
)

# --- 2段目：教科担当向け機能 ---
st.markdown('<div class="menu-group-header subject-header">📚 教科担当メイン機能</div>', unsafe_allow_html=True)
st.segmented_control(
    "教科担当機能",
    options=subject_options,
    default=st.session_state.active_menu if st.session_state.active_menu in subject_options else None,
    key="sub_menu",
    on_change=sync_menu_subject,
    label_visibility="collapsed"
)

st.divider()

selected_menu = st.session_state.active_menu

# ------------------------------------------
# 機能0: ⓪ 担任＆授業担当 名簿管理
# ------------------------------------------
if selected_menu == "⚙️ ⓪ 担任＆授業担当 名簿管理":
    st.subheader("⚙️ 全5クラス（200名）マスター名簿管理")
    m_tab1, m_tab2 = st.tabs(["🏫 担任＆全担当クラス名簿・新規追加", "📚 授業担当クラス名簿（コピペ・CSV出力）"])
    
    with m_tab1:
        col_m1, col_m2 = st.columns([1.3, 1])
        with col_m1:
            st.markdown("### 📋 全生徒マスター名簿（200名）")
            edited_master = st.data_editor(
                st.session_state.student_master,
                num_rows="dynamic",
                use_container_width=True,
                height=450,
                key="master_editor"
            )
            st.session_state.student_master = edited_master

        with col_m2:
            st.markdown("### 🔄 転入・転出手続き")
            action_type = st.radio("手続き種別:", ["生徒の新規登録・転入処理", "年度途中 転出（除籍）処理"])
            if action_type == "生徒の新規登録・転入処理":
                with st.form("trans_in_form"):
                    current_classes = get_all_classes()
                    in_cls_select = st.selectbox("登録クラス", current_classes + ["新規クラスを直接入力"])
                    in_cls_custom = st.text_input("新規クラス名（例: 2年3組）") if in_cls_select == "新規クラスを直接入力" else ""
                    target_cls = in_cls_custom if in_cls_select == "新規クラスを直接入力" else in_cls_select
                    in_num = st.number_input("出席番号", min_value=1, max_value=50, value=41)
                    in_name = st.text_input("生徒氏名")
                    in_gender = st.selectbox("性別", ["男", "女"])
                    in_date = st.date_input("登録日")
                    if st.form_submit_button("➕ 生徒を登録する") and in_name and target_cls:
                        new_st = pd.DataFrame([{"クラス": target_cls, "出席番号": in_num, "氏名": in_name, "性別": in_gender, "ステータス": "在籍", "異動日": str(in_date), "備考": "転入"}])
                        st.session_state.student_master = pd.concat([st.session_state.student_master, new_st], ignore_index=True)
                        st.success(f"{target_cls} に {in_name} さんを登録しました！")
                        st.rerun()

            else:
                with st.form("trans_out_form"):
                    active_list = get_active_students()
                    out_name = st.selectbox("転出生徒を選択", active_list)
                    out_date = st.date_input("転出日")
                    out_reason = st.text_input("転出理由")
                    if st.form_submit_button("⚠️ 転出処理を実行"):
                        idx = st.session_state.student_master[st.session_state.student_master["氏名"] == out_name].index
                        st.session_state.student_master.loc[idx, "ステータス"] = "転出"
                        st.session_state.student_master.loc[idx, "異動日"] = str(out_date)
                        st.session_state.student_master.loc[idx, "備考"] = out_reason
                        st.warning(f"{out_name} さんの転出処理を完了しました。")
                        st.rerun()

    with m_tab2:
        st.markdown("### 📚 担当5クラスの個別名簿抽出")
        all_classes = get_all_classes()
        selected_teach_cls = st.multiselect("抽出対象クラス選択:", all_classes, default=all_classes)
        if selected_teach_cls:
            sub_df = st.session_state.student_master[
                (st.session_state.student_master["クラス"].isin(selected_teach_cls)) &
                (st.session_state.student_master["ステータス"] == "在籍")
            ][["クラス", "出席番号", "氏名", "性別"]].sort_values(by=["クラス", "出席番号"])
            st.dataframe(sub_df, use_container_width=True, height=300)
            tsv_text = sub_df.to_csv(sep='\t', index=False)
            st.text_area("Excel貼り付け用タブ区切りテキスト (Ctrl+V用):", tsv_text, height=100)

# ------------------------------------------
# 機能1: ① 日々メモ蓄積 & 所見
# ------------------------------------------
elif selected_menu == "📝 ① 日々メモ蓄積 & 所見":
    col_a, col_b = st.columns([1, 1.2])
    all_cls = get_all_classes()
    
    with col_a:
        st.subheader("📌 観察メモを追加登録")
        with st.form("add_log_form", clear_on_submit=True):
            f_class = st.selectbox("クラス", all_cls, index=0)
            c_students = get_active_students(f_class)
            f_name = st.selectbox("生徒氏名", c_students)
            f_date = st.date_input("日付")
            f_cat = st.selectbox("対象分野", ["数学", "総合・行動の記録", "国語・他", "特別活動", "生活指導"])
            f_memo = st.text_area("観察メモ", placeholder="授業や学級活動での具体的な様子...")
            if st.form_submit_button("📥 メモをDBに保存") and f_name and f_memo:
                new_row = pd.DataFrame([{"日付": str(f_date), "クラス": f_class, "出席番号": 1, "氏名": f_name, "対象分野": f_cat, "観察メモ": f_memo}])
                st.session_state.daily_logs = pd.concat([st.session_state.daily_logs, new_row], ignore_index=True)
                st.success(f"{f_name} さんのメモを追加しました！")

    with col_b:
        st.subheader("✨ 蓄積メモからAI所見生成（担任1-1クラス）")
        c1_students = get_active_students("1年1組")
        selected_student = st.selectbox("所見を作成する1年1組の生徒:", c1_students, key="t1_st")
        
        student_memos = st.session_state.daily_logs[st.session_state.daily_logs["氏名"] == selected_student]
        st.write(f"📜 **{selected_student} さんの蓄積メモ（{len(student_memos)}件）:**")
        st.dataframe(student_memos[["日付", "対象分野", "観察メモ"]], use_container_width=True, height=200)
        
        if st.button("🪄 蓄積メモから所見文案を合成生成", type="primary"):
            if api_key and not student_memos.empty:
                combined_memos = "\n".join(student_memos["観察メモ"].tolist())
                prompt = f"生徒『{selected_student}』の蓄積メモ:\n{combined_memos}\n\n上記メモをもとに通知表用の所見文案を作成してください。{ending_rule}"
                model = genai.GenerativeModel("gemini-2.5-flash")
                with st.spinner("AIが所見文案を作成中..."):
                    res = model.generate_content(prompt)
                    st.text_area("生成された所見文案:", value=res.text.strip(), height=150)
            elif not api_key:
                st.warning("サイドバーでGemini API Keyを入力すると生成が試せます。")

# ------------------------------------------
# 機能2: ② CSV一括生成
# ------------------------------------------
elif selected_menu == "📁 ② CSV一括生成":
    st.subheader("📁 1年1組（40名）の蓄積メモから全員分所見をCSV一括出力")
    target_cls = st.selectbox("一括生成対象クラス:", ["1年1組"], key="t2_cls")
    
    if st.button("🚀 1年1組全員（40名分）の所見をAI一括生成"):
        cls_memos = st.session_state.daily_logs[st.session_state.daily_logs["クラス"] == target_cls]
        results = []
        
        if api_key:
            model = genai.GenerativeModel("gemini-2.5-flash")
            progress_bar = st.progress(0)
            unique_names = cls_memos["氏名"].unique()
            
            for i, name in enumerate(unique_names):
                group = cls_memos[cls_memos["氏名"] == name]
                all_memos = " / ".join(group["観察メモ"].tolist())
                res = model.generate_content(f"生徒:{name} メモ:{all_memos} の通知表所見を作成。{ending_rule}")
                results.append({"氏名": name, "まとめメモ": all_memos, "生成所見": res.text.strip()})
                progress_bar.progress((i + 1) / len(unique_names))
        else:
            for name, group in cls_memos.groupby("氏名"):
                all_memos = " / ".join(group["観察メモ"].tolist())
                results.append({"氏名": name, "まとめメモ": all_memos, "生成所見": "（APIキーを入力するとここにAI一括生成文が入ります）"})
                
        res_df = pd.DataFrame(results)
        st.dataframe(res_df, use_container_width=True, height=300)
        csv_data = res_df.to_csv(index=False).encode('utf-8-sig')
        st.download_button("📥 1年1組全員の所見一括CSVをダウンロード", csv_data, "1年1組_一括所見データ.csv", "text/csv")

# ------------------------------------------
# 機能3: ③ 所見自動校正
# ------------------------------------------
elif selected_menu == "🔍 ③ 所見データ自動校正":
    st.subheader("🔍 所見データの自動校正 & 不適切表現チェック")
    c1_students = get_active_students("1年1組")
    student_for_check = st.selectbox("校正を試す生徒を選択（1年1組）:", c1_students, key="chk_st")
    
    memos_text = " ".join(st.session_state.daily_logs[st.session_state.daily_logs["氏名"] == student_for_check]["観察メモ"].tolist())
    sample_text = st.text_area("校正対象テキスト:", value=f"{student_for_check}さんは、" + memos_text, height=120)
        
    if st.button("🛡️ AI誤字脱字・表現校正を実行", type="primary"):
        if api_key and sample_text:
            model = genai.GenerativeModel("gemini-2.5-flash")
            res = model.generate_content(f"誤字脱字チェックおよび保護者目線での適切な文章校正を行ってください:\n{sample_text}\n{ending_rule}")
            st.markdown("### 💡 校正結果アドバイス:")
            st.info(res.text)
        elif not api_key:
            st.warning("サイドバーでGemini API Keyを設定してください。")

# ------------------------------------------
# 機能4: ④ 面談用自動連携カルテ
# ------------------------------------------
elif selected_menu == "💬 ④ 蓄積連動カルテ":
    st.subheader("💬 保護者面談用カルテ（蓄積メモ＋テスト成績の自動連携）")
    kart_student = st.selectbox("面談対象生徒を選択（1年1組）:", get_active_students("1年1組"), key="kart_st")
    
    st_memos = st.session_state.daily_logs[st.session_state.daily_logs["氏名"] == kart_student]
    st_scores = st.session_state.subject_scores[st.session_state.subject_scores["氏名"] == kart_student]
    
    col_k1, col_k2 = st.columns(2)
    with col_k1:
        st.markdown(f"### 📜 日々の観察記録（{len(st_memos)}件）")
        st.dataframe(st_memos[["日付", "対象分野", "観察メモ"]], use_container_width=True, height=200)
    with col_k2:
        st.markdown("### 📊 自教科テスト＆観点成績")
        st.dataframe(st_scores[["中間テスト", "期末テスト", "観点1_知識", "観点2_思考", "観点3_主体性", "評定"]], use_container_width=True)
        
    if st.button("📋 面談用トークポイントカルテをAI生成", type="primary"):
        if api_key:
            model = genai.GenerativeModel("gemini-2.5-flash")
            memo_concat = " ".join(st_memos['観察メモ'].tolist())
            score_info = st_scores.to_dict(orient="records")[0] if not st_scores.empty else {}
            prompt = f"生徒『{kart_student}』の観察記録:{memo_concat}\nテスト成績:中間{score_info.get('中間テスト')}点, 期末{score_info.get('期末テスト')}点, 評定{score_info.get('評定')}\n面談で保護者に伝える【1.学習面・生活面の成長点 2.今後の課題 3.家庭での連携アドバイス】を簡潔に作成してください。"
            res = model.generate_content(prompt)
            st.info("💡 **AI生成 面談用カルテシート:**")
            st.markdown(res.text)
        else:
            st.warning("サイドバーでGemini API Keyを設定してください。")

# ------------------------------------------
# 機能5: ⑤ 全5クラス成績 & 評定自動計算
# ------------------------------------------
elif selected_menu == "📊 ⑤ 成績＆評定自動計算":
    st.subheader("📊 全5クラス（200名）成績・テスト得点・評定入力")
    sel_eval_cls = st.selectbox("表示クラス切替:", get_all_classes(), key="t5_cls")
    
    cls_score_df = st.session_state.subject_scores[st.session_state.subject_scores["クラス"] == sel_eval_cls]
    st.caption(f"※ {sel_eval_cls} の40名分のデータです。中間・期末テスト結果や観点を直接打ち替え可能です。")
    
    edited_scores = st.data_editor(cls_score_df, num_rows="dynamic", use_container_width=True, height=350, key=f"score_editor_{sel_eval_cls}")
    
    if st.button("⚡ 観点平均から5段階評定を一括全自動再計算"):
        def calc_grade(row):
            avg = (row["観点1_知識"] + row["観点2_思考"] + row["観点3_主体性"]) / 3
            if avg >= 85: return 5
            elif avg >= 70: return 4
            elif avg >= 55: return 3
            elif avg >= 40: return 2
            else: return 1
            
        st.session_state.subject_scores["評定"] = st.session_state.subject_scores.apply(calc_grade, axis=1)
        st.success("🎉 全5クラス（200名分）の観点データから5段階評定を再計算しました！")
        st.rerun()

# ------------------------------------------
# 機能6: ⑥ 学期推移ダッシュボード
# ------------------------------------------
elif selected_menu == "📈 ⑥ 学期推移ダッシュボード":
    st.subheader("📈 学期・テスト別 成績推移ダッシュボード")
    dash_cls = st.selectbox("ダッシュボード対象クラス:", get_all_classes(), key="t6_cls")
    
    cls_scores = st.session_state.subject_scores[st.session_state.subject_scores["クラス"] == dash_cls]
    
    st.markdown(f"### 📊 {dash_cls} 中間テスト vs 期末テスト 推移グラフ")
    chart_data = cls_scores.set_index("氏名")[["中間テスト", "期末テスト"]]
    st.line_chart(chart_data)
    
    st.markdown("### 📋 成績下降・フォロー対象者（期末テストで点数が下がった生徒）")
    down_students = cls_scores[cls_scores["期末テスト"] < cls_scores["中間テスト"]]
    st.dataframe(down_students[["出席番号", "氏名", "中間テスト", "期末テスト", "評定"]], use_container_width=True)

# ------------------------------------------
# 機能7: ⑦ 担任用 全教科成績集約
# ------------------------------------------
elif selected_menu == "🔄 ⑦ 担任用 全教科成績集約":
    st.subheader("🔄 各教科の成績ファイル自動名寄せ統合（ダミーデータ機能付き）")
    st.caption("他教科の担任から集まった個別CSVファイルを1人1行の全教科シートに合体します。")
    
    if st.button("🎲 他教科（国語・英語・理科・社会）の集約用ダミーデータを自動生成して結合テスト"):
        c1_students = st.session_state.student_master[st.session_state.student_master["クラス"] == "1年1組"]["氏名"].tolist()
        
        kokugo_df = pd.DataFrame([{"氏名": n, "国語_評定": random.randint(2, 5), "国語_観点_知識": random.randint(60, 95)} for n in c1_students])
        eigo_df = pd.DataFrame([{"氏名": n, "英語_評定": random.randint(2, 5), "英語_観点_知識": random.randint(55, 98)} for n in c1_students])
        
        math_df = st.session_state.subject_scores[st.session_state.subject_scores["クラス"] == "1年1組"][["氏名", "評定"]].rename(columns={"評定": "数学_評定"})
        
        merged_all = pd.merge(math_df, kokugo_df, on="氏名")
        merged_all = pd.merge(merged_all, eigo_df, on="氏名")
        
        st.success("🎉 1年1組の数学・国語・英語の成績データを『氏名』で完璧に合体しました！")
        st.dataframe(merged_all, use_container_width=True, height=350)
        
        csv_m = merged_all.to_csv(index=False).encode('utf-8-sig')
        st.download_button("📥 1年1組 全教科統合成績シート（CSV）を保存", csv_m, "1年1組_全教科成績統合表.csv", "text/csv")

# ------------------------------------------
# 機能8: 🖨️ 通知表・要録 印刷＆個別ファイル出力
# ------------------------------------------
elif selected_menu == "🖨️ ⑧ 通知表・要録 印刷＆出力":
    st.subheader("🖨️ 完成版・通知表プレビュー ＆ Wordダウンロード")
    
    col_out1, col_out2 = st.columns([1, 2])
    
    with col_out1:
        st.markdown("### 1. 対象生徒の選択")
        out_cls = st.selectbox("クラス選択:", get_all_classes(), index=0, key="t8_cls")
        out_students = get_active_students(out_cls)
        print_student = st.selectbox("生徒選択:", out_students)
        
        template_word = st.file_uploader("学校独自Wordテンプレート (.docx) の試用アップロード", type=["docx"])

    with col_out2:
        st.markdown(f"### 📄 画面プレビュー（{out_cls} {print_student} 様）")
        
        if print_student:
            st_info = st.session_state.student_master[st.session_state.student_master["氏名"] == print_student].iloc[0].to_dict()
            score_match = st.session_state.subject_scores[st.session_state.subject_scores["氏名"] == print_student]
            st_score = score_match.iloc[0].to_dict() if not score_match.empty else {}
            
            st.markdown(f"""
            <div class="student-card">
                <h3>【{st.session_state.school_year}年度 1学期 通知表プレビュー】</h3>
                <p><strong>所属:</strong> {st_info.get('クラス')} | <strong>出席番号:</strong> {st_info.get('出席番号')}番</p>
                <p><strong>氏名:</strong> <span style="font-size:1.3em; font-weight:bold;">{print_student}</span></p>
                <p><strong>学級担任:</strong> {st.session_state.teacher_name}</p>
                <hr>
                <h4>📊 自教科（数学）学習評価成績</h4>
                <ul>
                    <li>中間テスト: <strong>{st_score.get('中間テスト', '-')}</strong> 点 / 期末テスト: <strong>{st_score.get('期末テスト', '-')}</strong> 点</li>
                    <li>観点1 (知識・技能): <strong>{st_score.get('観点1_知識', '-')}</strong> 点</li>
                    <li>観点2 (思考・判断・表現): <strong>{st_score.get('観点2_思考', '-')}</strong> 点</li>
                    <li>観点3 (主体的に取り組む態度): <strong>{st_score.get('観点3_主体性', '-')}</strong> 点</li>
                    <li>学習評定（5段階）: <strong style="font-size:1.4em; color:#d9534f;">{st_score.get('評定', '-')}</strong></li>
                </ul>
                <hr>
                <h4>📝 総合所見（通知表文章）</h4>
                <p style="background-color: white; padding: 12px; border-radius: 6px; border: 1px solid #ccc; line-height: 1.6;">
                    {st_score.get('総合所見', '（所見データ準備完了）')}
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("📄 この生徒のWord通知表を出力・ダウンロード", type="primary"):
                doc_data = {
                    "年度": st.session_state.school_year,
                    "担任名": st.session_state.teacher_name,
                    "クラス": st_info.get('クラス'),
                    "出席番号": st_info.get('出席番号'),
                    "氏名": print_student,
                    "中間": st_score.get('中間テスト', ''),
                    "期末": st_score.get('期末テスト', ''),
                    "観点1": st_score.get('観点1_知識', ''),
                    "観点2": st_score.get('観点2_思考', ''),
                    "観点3": st_score.get('観点3_主体性', ''),
                    "評定": st_score.get('評定', ''),
                    "総合所見": st_score.get('総合所見', '')
                }
                
                if template_word:
                    doc = Document(template_word)
                else:
                    doc = Document()
                    doc.add_heading(f"通知表 - {print_student} 様", 0)
                    doc.add_paragraph(f"年度: {st.session_state.school_year}年度 | 担任: {st.session_state.teacher_name}")
                    doc.add_paragraph(f"クラス: {st_info.get('クラス')}  出席番号: {st_info.get('出席番号')}")
                    doc.add_paragraph(f"中間テスト: {st_score.get('中間テスト', '')}点 | 期末テスト: {st_score.get('期末テスト', '')}点")
                    doc.add_paragraph(f"数学評定: {st_score.get('評定', '')}")
                    doc.add_paragraph(f"総合所見:\n{st_score.get('総合所見', '')}")
                
                filled_doc = replace_docx_tags(doc, doc_data)
                out_buffer = io.BytesIO()
                filled_doc.save(out_buffer)
                
                st.download_button(
                    label=f"📥 {print_student} さんの Wordファイルを保存",
                    data=out_buffer.getvalue(),
                    file_name=f"通知表_{print_student}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )